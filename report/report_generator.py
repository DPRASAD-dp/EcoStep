import os
import sqlite3
import matplotlib.pyplot as plt
from datetime import datetime
from collections import defaultdict
from docx import Document
from docx.shared import Inches
from twilio.rest import Client
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from groq import Groq
import math

import config  # import from config.py

groq_client = Groq(api_key=config.GROQ_API_KEY)

class MonthlyReport:
    def __init__(self, month="2025-04", db_path="ocr_results.db"):
        self.month = month
        self.conn = sqlite3.connect(db_path, check_same_thread=False)
        self.cursor = self.conn.cursor()
        self.doc = Document()
        self.out_dir = f"report_{self.month}"
        os.makedirs(self.out_dir, exist_ok=True)

    def fetch_data(self):
        q = """SELECT item_name, carbon_footprint, category, current_date
               FROM receipts
               WHERE strftime('%Y-%m', current_date)=?"""
        return self.cursor.execute(q, (self.month,)).fetchall()

    def make_category_bar(self, data):
        cat_totals = defaultdict(float)
        for _, co2, cat, _ in data:
            cat_totals[cat] += co2

        f = os.path.join(self.out_dir, "cat_bar.png")
        if not cat_totals:
            plt.figure()
            plt.text(0.5, 0.5, "No data available", ha='center', va='center')
        else:
            sorted_items = sorted(cat_totals.items(), key=lambda x: x[1], reverse=True)
            cats, vals = zip(*sorted_items)

            plt.figure(figsize=(10, max(6, len(cats) * 0.5)))
            bars = plt.barh(cats, vals)
            for bar in bars:
                width = bar.get_width()
                plt.text(width + 0.1, bar.get_y() + bar.get_height()/2, f'{width:.2f}kg', ha='left', va='center')
            plt.title(f"CO₂ by Category ({self.month})")
            plt.xlabel("kg CO₂e")
            plt.tight_layout()

        plt.savefig(f, bbox_inches='tight')
        plt.close()
        return f

    def make_trend_chart(self, data):
        day_totals = defaultdict(float)
        for _, co2, _, dt in data:
            try:
                date_obj = datetime.strptime(dt, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                try:
                    date_obj = datetime.strptime(dt, "%Y-%m-%d")
                except ValueError:
                    print(f"Warning: Could not parse date: {dt}")
                    continue
            day = date_obj.day
            day_totals[day] += co2

        f = os.path.join(self.out_dir, "trend.png")
        if not day_totals:
            plt.figure()
            plt.text(0.5, 0.5, "No data available", ha='center', va='center')
        else:
            days = sorted(day_totals)
            vals = [day_totals[d] for d in days]

            plt.figure(figsize=(10, 6))
            plt.plot(days, vals, marker='o', linestyle='-', linewidth=2, markersize=8)
            for i, txt in enumerate(vals):
                plt.annotate(f"{txt:.2f}", (days[i], vals[i]), textcoords="offset points", xytext=(0,10), ha='center')
            plt.grid(True, linestyle='--', alpha=0.7)
            plt.title(f"Daily CO₂ Trend ({self.month})")
            plt.xlabel("Day of Month")
            plt.ylabel("kg CO₂e")
            plt.tight_layout()

        plt.savefig(f, bbox_inches='tight')
        plt.close()
        return f

    def make_category_pie(self, data):
        cat_totals = defaultdict(float)
        for _, co2, cat, _ in data:
            cat_totals[cat] += co2

        if not cat_totals:
            plt.figure()
            plt.text(0.5, 0.5, "No data available", ha='center', va='center')
            f = os.path.join(self.out_dir, "cat_pie.png")
            plt.savefig(f)
            plt.close()
            return f

        sorted_items = sorted(cat_totals.items(), key=lambda x: x[1], reverse=True)
        labels, sizes = zip(*sorted_items)

        plt.figure(figsize=(10, 8))
        wedges, _, autotexts = plt.pie(sizes, autopct='%1.1f%%', startangle=90)
        plt.setp(autotexts, size=10, weight="bold")
        plt.legend(wedges, labels, title="Categories", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
        plt.title(f"Category Share ({self.month})")
        plt.tight_layout()

        f = os.path.join(self.out_dir, "cat_pie.png")
        plt.savefig(f, bbox_inches='tight')
        plt.close()
        return f


    def summarize_text(self, data):
        lines = []
        for item, co2, cat, dt in data:
            try:
                date_obj = datetime.strptime(dt, "%Y-%m-%d %H:%M:%S")
            except ValueError:
                try:
                    date_obj = datetime.strptime(dt, "%Y-%m-%d")
                except ValueError:
                    date = dt
                else:
                    date = date_obj.strftime("%b %d")
            else:
                date = date_obj.strftime("%b %d")
            lines.append(f"{date} — {item}: {co2:.2f}kg CO₂e ({cat})")
        return "\n".join(lines) if lines else "No data available for this month."

    def get_suggestions(self, summary):
        if summary == "No data available for this month.":
            return "No data available to provide suggestions."
        prompt = f"Based on this monthly carbon footprint summary, give actionable suggestions to reduce CO₂:\n{summary}"
        try:
            response = groq_client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.3-70b-versatile",
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error getting suggestions: {str(e)}"

    def upload_to_drive(self, filepath):
        SCOPES = ['https://www.googleapis.com/auth/drive.file']
        flow = InstalledAppFlow.from_client_config({
            "installed": {
                "client_id": config.GOOGLE_CLIENT_ID,
                "client_secret": config.GOOGLE_CLIENT_SECRET,
                "auth_uri": "https://accounts.google.com/o/oauth2/auth",
                "token_uri": "https://oauth2.googleapis.com/token",
                "redirect_uris": ["http://localhost"]
            }
        }, SCOPES)

        creds = flow.run_local_server(port=0)
        service = build('drive', 'v3', credentials=creds)

        file_metadata = {
            'name': os.path.basename(filepath),
            'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }
        media = MediaFileUpload(filepath, mimetype=file_metadata['mimeType'])
        file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()

        service.permissions().create(
            fileId=file['id'],
            body={'type': 'anyone', 'role': 'reader'},
        ).execute()

        return f"https://drive.google.com/file/d/{file['id']}/view?usp=sharing"

    def send_whatsapp(self, message):
        try:
            client = Client(config.TWILIO_ACCOUNT_SID, config.TWILIO_AUTH_TOKEN)
            client.messages.create(body=message, from_=config.WHATSAPP_FROM, to=config.WHATSAPP_TO)
            print("✅ WhatsApp message sent.")
        except Exception as e:
            print(f"❌ WhatsApp send failed: {e}")

    def send_notifications(self, link, summary):
        message = f"📄 Monthly Carbon Footprint Report ({self.month})\n\n📥 Download Report: {link}"
        self.send_whatsapp(message)

    def build(self):
        data = self.fetch_data()
        out_file = f"Monthly_Report_{self.month}.docx"

        self.doc.add_heading(f"Carbon Footprint Report — {self.month}", level=1)
        if not data:
            self.doc.add_paragraph("No carbon footprint data available for this month.")
            self.doc.save(out_file)
            print("Saved empty report to", out_file)
            link = self.upload_to_drive(out_file)
            self.send_notifications(link, "No data available for this month.")
            return

        self.doc.add_picture(self.make_category_bar(data), width=Inches(6))
        self.doc.add_picture(self.make_trend_chart(data), width=Inches(6))
        self.doc.add_picture(self.make_category_pie(data), width=Inches(6))

        summary = self.summarize_text(data)
        self.doc.add_heading("Summary", level=2)
        self.doc.add_paragraph(summary)

        suggestions = self.get_suggestions(summary)
        self.doc.add_heading("AI Suggestions", level=2)
        self.doc.add_paragraph(suggestions)

        self.doc.save(out_file)
        print("✅ Report saved to", out_file)

        link = self.upload_to_drive(out_file)
        self.send_notifications(link, summary)

if __name__ == "__main__":
    MonthlyReport().build()